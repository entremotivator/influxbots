#!/usr/bin/env python3
"""
🚀 ADVANCED LANGCHAIN STREAMLIT CHAT BOT
Comprehensive AI platform with document processing, vector search, memory management, and intelligent agents
Optimized for Streamlit Cloud deployment with extensive LangChain ecosystem integration
"""

import streamlit as st
import os
import json
import time
import hashlib
import pandas as pd
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Optional, Any, Union
import logging
from io import StringIO, BytesIO
import base64
import tempfile
import asyncio
from pathlib import Path
import uuid

# Core imports
try:
    import openai
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

# LangChain Core imports
try:
    from langchain.llms import OpenAI as LangChainOpenAI
    from langchain.chat_models import ChatOpenAI
    from langchain.schema import Document, BaseMessage, HumanMessage, AIMessage, SystemMessage
    LANGCHAIN_CORE_AVAILABLE = True
except ImportError:
    LANGCHAIN_CORE_AVAILABLE = False

# Additional processing libraries
try:
    import PyPDF2
    from PIL import Image
    import docx
    from openpyxl import load_workbook
    import pandas as pd
    import numpy as np
    FILE_PROCESSING_AVAILABLE = True
except ImportError:
    FILE_PROCESSING_AVAILABLE = False

# Advanced processing libraries
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

# Web scraping and API libraries
try:
    import requests
    from bs4 import BeautifulSoup
    WEB_PROCESSING_AVAILABLE = True
except ImportError:
    WEB_PROCESSING_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ======================================================
# 🔐 CONFIGURATION MANAGEMENT
# ======================================================

class ConfigurationManager:
    """Configuration management for LangChain Streamlit app"""
    
    def __init__(self):
        self.config = self.load_configuration()
    
    def load_configuration(self) -> Dict[str, Any]:
        """Load configuration from multiple sources"""
        try:
            config = {
                # OpenAI Configuration
                "openai_api_key": self._get_config_value("OPENAI_API_KEY", ""),
                "openai_model": self._get_config_value("OPENAI_MODEL", "gpt-4"),
                "openai_temperature": float(self._get_config_value("OPENAI_TEMPERATURE", "0.7")),
                "openai_max_tokens": int(self._get_config_value("OPENAI_MAX_TOKENS", "2000")),
                
                # File Processing Configuration
                "max_file_size_mb": int(self._get_config_value("MAX_FILE_SIZE_MB", "100")),
                "max_files_per_session": int(self._get_config_value("MAX_FILES_PER_SESSION", "50")),
                "supported_file_types": self._get_config_value("SUPPORTED_FILE_TYPES", "pdf,docx,txt,csv,json,xlsx,html,md,py,js").split(","),
                
                # Application Configuration
                "app_name": self._get_config_value("APP_NAME", "Advanced LangChain AI Assistant"),
                "app_description": self._get_config_value("APP_DESCRIPTION", "Comprehensive AI platform with document processing"),
                "debug_mode": self._get_config_value("DEBUG_MODE", "false").lower() == "true",
                
                # UI Configuration
                "theme_primary_color": self._get_config_value("THEME_PRIMARY_COLOR", "#667eea"),
                "theme_secondary_color": self._get_config_value("THEME_SECONDARY_COLOR", "#764ba2"),
            }
            
            return config
            
        except Exception as e:
            logger.error(f"Error loading configuration: {str(e)}")
            return self._get_default_configuration()
    
    def _get_config_value(self, key: str, default: str) -> str:
        """Get configuration value from Streamlit secrets or environment variables"""
        try:
            # Try Streamlit secrets first
            if hasattr(st, 'secrets') and key in st.secrets:
                return st.secrets[key]
            
            # Fall back to environment variables
            return os.getenv(key, default)
            
        except Exception:
            return default
    
    def _get_default_configuration(self) -> Dict[str, Any]:
        """Get default configuration when loading fails"""
        return {
            "openai_api_key": "",
            "openai_model": "gpt-4",
            "openai_temperature": 0.7,
            "openai_max_tokens": 2000,
            "max_file_size_mb": 100,
            "max_files_per_session": 50,
            "supported_file_types": ["pdf", "docx", "txt", "csv", "json", "xlsx", "html", "md", "py", "js"],
            "app_name": "Advanced LangChain AI Assistant",
            "app_description": "Comprehensive AI platform with document processing",
            "debug_mode": False,
            "theme_primary_color": "#667eea",
            "theme_secondary_color": "#764ba2"
        }
    
    def get(self, key: str, default: Any = None) -> Any:
        """Get configuration value"""
        return self.config.get(key, default)
    
    def update(self, key: str, value: Any):
        """Update configuration value"""
        self.config[key] = value

# Initialize configuration manager
config_manager = ConfigurationManager()

# ======================================================
# 🤖 AI ASSISTANT PROFILES
# ======================================================

class AIAssistantProfiles:
    """AI assistant profiles with specialized capabilities"""
    
    @staticmethod
    def get_all_assistants() -> Dict[str, Dict[str, Any]]:
        """Get all available AI assistants"""
        return {
            "Strategic Business Consultant": {
                "description": "Senior strategic consultant with expertise in business transformation and growth strategy.",
                "system_prompt": "You are a strategic business consultant with 20+ years of experience. Provide data-driven insights, actionable recommendations, and strategic frameworks. Focus on both short-term tactics and long-term strategic implications.",
                "emoji": "🎯",
                "category": "Business Strategy",
                "specialties": ["Strategic Planning", "Market Analysis", "Business Transformation"],
                "expertise_level": "Senior Partner",
                "temperature": 0.3,
                "max_tokens": 3000
            },
            
            "Startup Growth Advisor": {
                "description": "Experienced startup mentor specializing in scaling early-stage companies.",
                "system_prompt": "You are a startup growth advisor who has mentored 200+ startups. Be practical, action-oriented, and focused on measurable results. Provide specific tactics for resource-constrained environments.",
                "emoji": "🚀",
                "category": "Entrepreneurship",
                "specialties": ["Product-Market Fit", "Growth Hacking", "Fundraising"],
                "expertise_level": "Serial Entrepreneur",
                "temperature": 0.7,
                "max_tokens": 2500
            },
            
            "Digital Marketing Strategist": {
                "description": "Performance marketing expert driving growth through data-driven campaigns.",
                "system_prompt": "You are a digital marketing strategist focused on performance marketing and growth analytics. Provide measurable ROI strategies, specific metrics, and testing frameworks.",
                "emoji": "📈",
                "category": "Marketing & Growth",
                "specialties": ["Performance Marketing", "Customer Acquisition", "Conversion Optimization"],
                "expertise_level": "VP Marketing",
                "temperature": 0.4,
                "max_tokens": 2800
            },
            
            "Sales Performance Coach": {
                "description": "Elite sales trainer helping maximize revenue through proven methodologies.",
                "system_prompt": "You are a sales performance coach with expertise in consultative selling and sales psychology. Provide practical, actionable sales techniques focused on long-term relationships.",
                "emoji": "💰",
                "category": "Sales & Revenue",
                "specialties": ["Consultative Selling", "Objection Handling", "Pipeline Management"],
                "expertise_level": "VP Sales",
                "temperature": 0.5,
                "max_tokens": 2500
            },
            
            "Financial Strategy Advisor": {
                "description": "CFO-level financial expert optimizing business finances and strategic planning.",
                "system_prompt": "You are a financial strategy advisor with CFO-level expertise. Provide detailed financial analysis with calculations, focus on sustainable growth and risk mitigation.",
                "emoji": "💼",
                "category": "Finance & Investment",
                "specialties": ["Financial Planning", "Investment Strategy", "Risk Management"],
                "expertise_level": "CFO",
                "temperature": 0.2,
                "max_tokens": 3500
            },
            
            "Operations Excellence Manager": {
                "description": "Process optimization expert streamlining operations through lean methodologies.",
                "system_prompt": "You are an operations excellence manager with Lean Six Sigma expertise. Think systematically about processes, identify bottlenecks, and provide scalable improvement solutions.",
                "emoji": "⚙️",
                "category": "Operations & Process",
                "specialties": ["Process Optimization", "Lean Six Sigma", "Quality Management"],
                "expertise_level": "VP Operations",
                "temperature": 0.3,
                "max_tokens": 3000
            },
            
            "Technology Strategy Consultant": {
                "description": "CTO-level technology advisor guiding digital transformation and innovation.",
                "system_prompt": "You are a technology strategy consultant with CTO-level expertise. Balance technical feasibility with business value, provide implementation roadmaps with risk assessments.",
                "emoji": "💻",
                "category": "Technology & Innovation",
                "specialties": ["Digital Transformation", "Technology Architecture", "Innovation Strategy"],
                "expertise_level": "CTO",
                "temperature": 0.4,
                "max_tokens": 3200
            },
            
            "AI & Machine Learning Specialist": {
                "description": "Advanced AI researcher helping organizations leverage artificial intelligence.",
                "system_prompt": "You are an AI and machine learning specialist with deep expertise. Provide technically accurate guidance while making concepts accessible. Focus on practical applications and business value.",
                "emoji": "🤖",
                "category": "Artificial Intelligence",
                "specialties": ["Machine Learning", "Natural Language Processing", "Data Science"],
                "expertise_level": "AI Research Director",
                "temperature": 0.3,
                "max_tokens": 3500
            },
            
            "Human Capital Strategist": {
                "description": "CHRO-level HR expert focused on talent strategy and organizational development.",
                "system_prompt": "You are a human capital strategist with CHRO-level expertise. Focus on people-centric solutions that drive business results. Consider both individual and organizational perspectives.",
                "emoji": "👥",
                "category": "Human Resources",
                "specialties": ["Talent Strategy", "Organizational Development", "Culture Transformation"],
                "expertise_level": "CHRO",
                "temperature": 0.6,
                "max_tokens": 2800
            },
            
            "Customer Experience Director": {
                "description": "Customer-centric leader designing exceptional experiences for business growth.",
                "system_prompt": "You are a customer experience director focused on creating exceptional customer journeys. Always think from the customer's perspective first and focus on emotional connections.",
                "emoji": "🤝",
                "category": "Customer Success",
                "specialties": ["Customer Journey Design", "Experience Strategy", "Customer Retention"],
                "expertise_level": "VP Customer Experience",
                "temperature": 0.5,
                "max_tokens": 2700
            },
            
            "Product Management Expert": {
                "description": "Senior product leader with expertise in product strategy and user-centered design.",
                "system_prompt": "You are a product management expert with extensive experience in product strategy. Focus on user needs and business value with data-driven decisions and clear prioritization.",
                "emoji": "📱",
                "category": "Product Management",
                "specialties": ["Product Strategy", "User Research", "Product-Market Fit"],
                "expertise_level": "VP Product",
                "temperature": 0.4,
                "max_tokens": 2900
            },
            
            "Data Science Consultant": {
                "description": "Senior data scientist turning complex data into actionable business insights.",
                "system_prompt": "You are a senior data scientist with expertise in advanced analytics. Provide statistically sound analysis with clear business implications. Make complex analyses accessible through clear explanations.",
                "emoji": "📊",
                "category": "Data & Analytics",
                "specialties": ["Statistical Analysis", "Predictive Modeling", "Data Visualization"],
                "expertise_level": "Senior Data Scientist",
                "temperature": 0.2,
                "max_tokens": 3200
            },
            
            "Document Intelligence Specialist": {
                "description": "Advanced document processing expert specializing in information extraction.",
                "system_prompt": "You are a document intelligence specialist with expertise in document processing and knowledge management. Provide comprehensive analysis with structured insights and actionable information.",
                "emoji": "📄",
                "category": "Document Intelligence",
                "specialties": ["Document Analysis", "Information Extraction", "Knowledge Management"],
                "expertise_level": "Document Intelligence Expert",
                "temperature": 0.3,
                "max_tokens": 3500
            },
            
            "Content Strategy Director": {
                "description": "Content marketing leader creating engaging, conversion-focused content strategies.",
                "system_prompt": "You are a content strategy director with expertise in content marketing and brand storytelling. Focus on creating content that drives business results and audience engagement.",
                "emoji": "✍️",
                "category": "Content & Communication",
                "specialties": ["Content Strategy", "Brand Storytelling", "Multi-channel Distribution"],
                "expertise_level": "VP Content",
                "temperature": 0.6,
                "max_tokens": 2800
            },
            
            "Agile Project Manager": {
                "description": "Certified project management expert ensuring successful delivery through agile methodologies.",
                "system_prompt": "You are an agile project manager with PMP and Scrum Master certifications. Provide structured frameworks with clear deliverables, timelines, and risk mitigation strategies.",
                "emoji": "📋",
                "category": "Project Management",
                "specialties": ["Agile Methodology", "Risk Management", "Stakeholder Management"],
                "expertise_level": "Senior Project Manager",
                "temperature": 0.4,
                "max_tokens": 2600
            }
        }

# Initialize assistant profiles
assistant_profiles = AIAssistantProfiles()

# ======================================================
# 📁 ADVANCED FILE PROCESSOR
# ======================================================

class AdvancedFileProcessor:
    """Advanced file processing with comprehensive analysis"""
    
    def __init__(self, config_manager: ConfigurationManager):
        self.config = config_manager
        self.processed_files = {}
        self.temp_dir = tempfile.mkdtemp()
        
    def process_uploaded_files(self, uploaded_files: List[Any], session_id: str) -> Dict[str, Any]:
        """Process multiple uploaded files with comprehensive analysis"""
        try:
            if not uploaded_files:
                return {"status": "error", "message": "No files provided"}
            
            # Validate file limits
            if len(uploaded_files) > self.config.get("max_files_per_session", 50):
                return {
                    "status": "error", 
                    "message": f"Too many files. Maximum {self.config.get('max_files_per_session')} files per session."
                }
            
            processing_results = {
                "session_id": session_id,
                "total_files": len(uploaded_files),
                "processed_files": [],
                "failed_files": [],
                "summary": "",
                "processing_time": 0,
                "status": "success"
            }
            
            start_time = time.time()
            
            # Process each file
            for uploaded_file in uploaded_files:
                try:
                    # Save uploaded file to temporary location
                    temp_path = os.path.join(self.temp_dir, f"{session_id}_{uploaded_file.name}")
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.read())
                    
                    # Process individual file
                    file_result = self.process_single_file(uploaded_file, temp_path)
                    
                    if file_result["status"] == "success":
                        processing_results["processed_files"].append(file_result)
                    else:
                        processing_results["failed_files"].append(file_result)
                        
                except Exception as e:
                    logger.error(f"Error processing file {uploaded_file.name}: {str(e)}")
                    processing_results["failed_files"].append({
                        "name": uploaded_file.name,
                        "error": str(e),
                        "status": "error"
                    })
            
            # Generate comprehensive summary
            processing_results["summary"] = self.generate_processing_summary(processing_results)
            processing_results["processing_time"] = time.time() - start_time
            
            # Store results
            self.processed_files[session_id] = processing_results
            
            return processing_results
            
        except Exception as e:
            logger.error(f"Error in batch file processing: {str(e)}")
            return {
                "status": "error",
                "message": str(e),
                "session_id": session_id
            }
    
    def process_single_file(self, uploaded_file: Any, temp_path: str) -> Dict[str, Any]:
        """Process a single file with comprehensive analysis"""
        try:
            file_info = {
                "name": uploaded_file.name,
                "size": uploaded_file.size,
                "type": uploaded_file.type,
                "file_extension": Path(uploaded_file.name).suffix.lower(),
                "processing_method": "unknown",
                "content": "",
                "metadata": {},
                "summary": "",
                "key_insights": [],
                "status": "success",
                "processing_time": 0,
                "error": None
            }
            
            start_time = time.time()
            
            # Validate file
            validation_result = self.validate_file(uploaded_file)
            if not validation_result["valid"]:
                file_info["status"] = "error"
                file_info["error"] = validation_result["message"]
                return file_info
            
            # Determine processing method
            file_extension = file_info["file_extension"].lstrip('.')
            
            # Process based on file type
            if file_extension == "pdf":
                file_info.update(self.process_pdf_custom(uploaded_file))
            elif file_extension in ["docx", "doc"]:
                file_info.update(self.process_docx_custom(uploaded_file))
            elif file_extension == "csv":
                file_info.update(self.process_csv_custom(uploaded_file))
            elif file_extension in ["xlsx", "xls"]:
                file_info.update(self.process_excel_custom(uploaded_file))
            elif file_extension == "json":
                file_info.update(self.process_json_custom(uploaded_file))
            elif file_extension in ["txt", "md", "py", "js", "html", "css", "sql"]:
                file_info.update(self.process_text_custom(uploaded_file))
            elif file_extension in ["jpg", "jpeg", "png", "gif", "bmp", "tiff"]:
                file_info.update(self.process_image_custom(uploaded_file))
            else:
                file_info.update({
                    "processing_method": "unsupported",
                    "content": f"Unsupported file type: .{file_extension}",
                    "error": f"No processor available for .{file_extension} files"
                })
            
            # Generate insights and summary
            if file_info["content"]:
                file_info["summary"] = self.generate_file_summary(file_info)
                file_info["key_insights"] = self.extract_key_insights(file_info)
            
            file_info["processing_time"] = time.time() - start_time
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error processing file {uploaded_file.name}: {str(e)}")
            return {
                "name": uploaded_file.name,
                "status": "error",
                "error": str(e),
                "processing_time": time.time() - start_time if 'start_time' in locals() else 0
            }
    
    def validate_file(self, uploaded_file: Any) -> Dict[str, Any]:
        """Validate uploaded file"""
        try:
            max_size = self.config.get("max_file_size_mb", 100) * 1024 * 1024
            supported_types = self.config.get("supported_file_types", [])
            
            if uploaded_file.size > max_size:
                return {
                    "valid": False,
                    "message": f"File too large. Maximum size: {self.config.get('max_file_size_mb')}MB"
                }
            
            file_extension = Path(uploaded_file.name).suffix.lower().lstrip('.')
            if supported_types and file_extension not in supported_types:
                return {
                    "valid": False,
                    "message": f"Unsupported file type: .{file_extension}"
                }
            
            return {"valid": True, "message": "File is valid"}
            
        except Exception as e:
            return {"valid": False, "message": f"Validation error: {str(e)}"}
    
    def process_pdf_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom PDF processing"""
        try:
            if not FILE_PROCESSING_AVAILABLE:
                return {"processing_method": "error", "content": "", "error": "PDF processing not available"}
            
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = []
            metadata = {
                "total_pages": len(pdf_reader.pages),
                "pdf_info": pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
            }
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"=== PAGE {page_num + 1} ===\n{text}")
                except Exception as e:
                    text_content.append(f"=== PAGE {page_num + 1} ===\nError extracting text: {str(e)}")
            
            content = "\n\n".join(text_content) if text_content else "No text content extracted from PDF"
            
            return {
                "processing_method": "custom_pdf",
                "content": content,
                "metadata": metadata
            }
            
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"PDF processing error: {str(e)}"}
    
    def process_docx_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom DOCX processing"""
        try:
            if not FILE_PROCESSING_AVAILABLE:
                return {"processing_method": "error", "content": "", "error": "DOCX processing not available"}
            
            doc = docx.Document(uploaded_file)
            text_content = []
            metadata = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables)
            }
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("=== TABLE ===\n" + "\n".join(table_text))
            
            content = "\n\n".join(text_content) if text_content else "No text content found in document"
            
            return {
                "processing_method": "custom_docx",
                "content": content,
                "metadata": metadata
            }
            
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"DOCX processing error: {str(e)}"}
    
    def process_csv_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom CSV processing with advanced analysis"""
        try:
            if not FILE_PROCESSING_AVAILABLE:
                return {"processing_method": "error", "content": "", "error": "CSV processing not available"}
            
            # Read CSV content
            content = uploaded_file.read().decode('utf-8')
            uploaded_file.seek(0)
            
            df = pd.read_csv(StringIO(content))
            
            # Generate comprehensive analysis
            analysis = f"=== CSV FILE ANALYSIS ===\n\n"
            analysis += f"📊 **Dataset Overview:**\n"
            analysis += f"• Rows: {len(df):,}\n"
            analysis += f"• Columns: {len(df.columns)}\n"
            analysis += f"• Total cells: {df.size:,}\n"
            
            analysis += f"\n📋 **Column Information:**\n"
            for col in df.columns:
                dtype = str(df[col].dtype)
                null_count = df[col].isnull().sum()
                unique_count = df[col].nunique()
                analysis += f"• {col}: {dtype} ({unique_count:,} unique, {null_count:,} null)\n"
            
            # Sample data
            analysis += f"\n📋 **Sample Data (first 5 rows):**\n"
            sample_data = df.head().to_string(index=False)
            analysis += sample_data
            
            if len(df) > 5:
                analysis += f"\n... and {len(df) - 5:,} more rows"
            
            # Add raw content for analysis
            analysis += f"\n\n=== RAW CSV CONTENT ===\n{content[:2000]}"
            if len(content) > 2000:
                analysis += "... [Content truncated for processing]"
            
            metadata = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "data_types": df.dtypes.to_dict()
            }
            
            return {
                "processing_method": "custom_csv_advanced",
                "content": analysis,
                "metadata": metadata
            }
            
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"CSV processing error: {str(e)}"}
    
    def process_excel_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom Excel processing"""
        try:
            if not FILE_PROCESSING_AVAILABLE:
                return {"processing_method": "error", "content": "", "error": "Excel processing not available"}
            
            workbook = load_workbook(uploaded_file, data_only=True)
            
            analysis = f"=== EXCEL FILE ANALYSIS ===\n\n"
            analysis += f"📊 **Workbook Overview:**\n"
            analysis += f"• Total worksheets: {len(workbook.sheetnames)}\n"
            analysis += f"• Sheet names: {', '.join(workbook.sheetnames)}\n\n"
            
            metadata = {
                "total_sheets": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "sheets_data": {}
            }
            
            # Process each sheet (limit to first 3 sheets)
            for sheet_name in workbook.sheetnames[:3]:
                sheet = workbook[sheet_name]
                
                analysis += f"📋 **Sheet: '{sheet_name}'**\n"
                analysis += f"• Dimensions: {sheet.max_row} rows × {sheet.max_column} columns\n"
                
                # Extract sample data
                sheet_data = []
                for row_num in range(1, min(6, sheet.max_row + 1)):  # First 5 rows
                    row_data = []
                    for col_num in range(1, min(6, sheet.max_column + 1)):  # First 5 columns
                        cell_value = sheet.cell(row=row_num, column=col_num).value
                        row_data.append(str(cell_value) if cell_value is not None else "")
                    sheet_data.append(" | ".join(row_data))
                
                if sheet_data:
                    analysis += f"Sample data:\n"
                    for i, row in enumerate(sheet_data):
                        analysis += f"Row {i+1}: {row}\n"
                
                analysis += "\n"
                
                metadata["sheets_data"][sheet_name] = {
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column
                }
            
            return {
                "processing_method": "custom_excel",
                "content": analysis,
                "metadata": metadata
            }
            
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"Excel processing error: {str(e)}"}
    
    def process_json_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom JSON processing"""
        try:
            content = uploaded_file.read().decode('utf-8')
            data = json.loads(content)
            
            analysis = f"=== JSON FILE ANALYSIS ===\n\n"
            analysis += f"📊 **Structure Overview:**\n"
            analysis += f"• Root type: {type(data).__name__}\n"
            analysis += f"• File size: {len(content):,} characters\n"
            
            analysis += f"\n📋 **Content Preview:**\n"
            preview = json.dumps(data, indent=2)[:1000]
            analysis += preview
            if len(json.dumps(data, indent=2)) > 1000:
                analysis += "\n... [Content truncated for processing]"
            
            metadata = {
                "root_type": type(data).__name__,
                "file_size_chars": len(content)
            }
            
            if isinstance(data, dict):
                metadata["top_level_keys"] = list(data.keys())[:10]  # First 10 keys
            elif isinstance(data, list):
                metadata["array_length"] = len(data)
            
            return {
                "processing_method": "custom_json",
                "content": analysis,
                "metadata": metadata
            }
            
        except json.JSONDecodeError as e:
            return {"processing_method": "error", "content": "", "error": f"Invalid JSON format: {str(e)}"}
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"JSON processing error: {str(e)}"}
    
    def process_text_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom text file processing"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    content = uploaded_file.read().decode(encoding)
                    uploaded_file.seek(0)
                    break
                except UnicodeDecodeError:
                    uploaded_file.seek(0)
                    continue
            
            if content is None:
                return {"processing_method": "error", "content": "", "error": "Could not decode file"}
            
            # Analyze text content
            lines = content.split('\n')
            words = content.split()
            
            analysis = f"=== TEXT FILE ANALYSIS ===\n\n"
            analysis += f"📊 **Content Statistics:**\n"
            analysis += f"• Total lines: {len(lines):,}\n"
            analysis += f"• Total words: {len(words):,}\n"
            analysis += f"• Total characters: {len(content):,}\n"
            
            analysis += f"\n📋 **Content Preview (first 10 lines):**\n"
            preview_lines = lines[:10]
            for i, line in enumerate(preview_lines, 1):
                analysis += f"{i:3d}: {line}\n"
            
            if len(lines) > 10:
                analysis += f"... and {len(lines) - 10:,} more lines\n"
            
            analysis += f"\n\n=== FULL CONTENT ===\n{content}"
            
            metadata = {
                "total_lines": len(lines),
                "total_words": len(words),
                "total_characters": len(content)
            }
            
            return {
                "processing_method": "custom_text",
                "content": analysis,
                "metadata": metadata
            }
            
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"Text processing error: {str(e)}"}
    
    def process_image_custom(self, uploaded_file: Any) -> Dict[str, Any]:
        """Custom image processing"""
        try:
            if not FILE_PROCESSING_AVAILABLE:
                return {"processing_method": "error", "content": "", "error": "Image processing not available"}
            
            image = Image.open(uploaded_file)
            
            analysis = f"=== IMAGE FILE ANALYSIS ===\n\n"
            analysis += f"🖼️ **Image Properties:**\n"
            analysis += f"• Format: {image.format}\n"
            analysis += f"• Dimensions: {image.size[0]} × {image.size[1]} pixels\n"
            analysis += f"• Color mode: {image.mode}\n"
            analysis += f"• File size: {uploaded_file.size:,} bytes\n"
            
            # Calculate additional properties
            width, height = image.size
            aspect_ratio = width / height
            megapixels = (width * height) / 1_000_000
            
            analysis += f"• Aspect ratio: {aspect_ratio:.2f}:1\n"
            analysis += f"• Megapixels: {megapixels:.1f}MP\n"
            
            analysis += f"\n📝 **Processing Notes:**\n"
            analysis += f"• Image successfully loaded and analyzed\n"
            analysis += f"• Ready for AI-powered visual analysis\n"
            
            metadata = {
                "format": image.format,
                "width": width,
                "height": height,
                "mode": image.mode,
                "file_size": uploaded_file.size,
                "aspect_ratio": aspect_ratio,
                "megapixels": megapixels
            }
            
            return {
                "processing_method": "custom_image",
                "content": analysis,
                "metadata": metadata
            }
            
        except Exception as e:
            return {"processing_method": "error", "content": "", "error": f"Image processing error: {str(e)}"}
    
    def generate_file_summary(self, file_info: Dict[str, Any]) -> str:
        """Generate a comprehensive summary of the processed file"""
        try:
            name = file_info.get("name", "Unknown")
            size = file_info.get("size", 0)
            processing_method = file_info.get("processing_method", "unknown")
            content_length = len(file_info.get("content", ""))
            
            summary = f"📄 **{name}**\n"
            summary += f"• Size: {size:,} bytes\n"
            summary += f"• Processing: {processing_method}\n"
            summary += f"• Content extracted: {content_length:,} characters\n"
            
            processing_time = file_info.get("processing_time", 0)
            summary += f"• Processing time: {processing_time:.2f}s\n"
            
            return summary
            
        except Exception as e:
            return f"Summary generation failed: {str(e)}"
    
    def extract_key_insights(self, file_info: Dict[str, Any]) -> List[str]:
        """Extract key insights from processed file"""
        try:
            insights = []
            content = file_info.get("content", "")
            metadata = file_info.get("metadata", {})
            file_extension = file_info.get("file_extension", "").lstrip('.')
            
            # File type specific insights
            if file_extension == "csv":
                if "rows" in metadata and "columns" in metadata:
                    rows, cols = metadata["rows"], metadata["columns"]
                    if rows > 1000:
                        insights.append(f"Large dataset with {rows:,} rows - suitable for analysis")
                    if cols > 10:
                        insights.append(f"Wide dataset with {cols} columns")
            
            # Content-based insights
            if content:
                word_count = len(content.split())
                if word_count > 1000:
                    insights.append("Substantial content - suitable for detailed analysis")
                elif word_count < 100:
                    insights.append("Brief content - may need additional context")
            
            return insights[:3]  # Limit to top 3 insights
            
        except Exception as e:
            logger.error(f"Error extracting insights: {str(e)}")
            return []
    
    def generate_processing_summary(self, processing_results: Dict[str, Any]) -> str:
        """Generate comprehensive summary of batch processing results"""
        try:
            total_files = processing_results.get("total_files", 0)
            processed_files = processing_results.get("processed_files", [])
            failed_files = processing_results.get("failed_files", [])
            processing_time = processing_results.get("processing_time", 0)
            
            summary = f"📊 **Batch Processing Summary**\n\n"
            summary += f"• Total files: {total_files}\n"
            summary += f"• Successfully processed: {len(processed_files)}\n"
            summary += f"• Failed: {len(failed_files)}\n"
            summary += f"• Processing time: {processing_time:.2f}s\n"
            
            if processed_files:
                summary += f"\n✅ **Successfully Processed:**\n"
                for file_info in processed_files[:5]:  # Show first 5
                    name = file_info.get("name", "Unknown")
                    method = file_info.get("processing_method", "unknown")
                    summary += f"• {name} ({method})\n"
                
                if len(processed_files) > 5:
                    summary += f"• ... and {len(processed_files) - 5} more files\n"
            
            if failed_files:
                summary += f"\n❌ **Failed to Process:**\n"
                for file_info in failed_files[:3]:  # Show first 3 failures
                    name = file_info.get("name", "Unknown")
                    error = file_info.get("error", "Unknown error")
                    summary += f"• {name}: {error}\n"
            
            return summary
            
        except Exception as e:
            return f"Summary generation failed: {str(e)}"

# Initialize advanced file processor
advanced_file_processor = AdvancedFileProcessor(config_manager)

# ======================================================
# 🤖 AI CHAT MANAGER
# ======================================================

class AIChatManager:
    """AI chat manager with OpenAI integration"""
    
    def __init__(self, config_manager: ConfigurationManager):
        self.config = config_manager
        self.client = None
        self.initialize_client()
    
    def initialize_client(self):
        """Initialize OpenAI client"""
        try:
            api_key = self.config.get("openai_api_key")
            if api_key and OPENAI_AVAILABLE:
                self.client = OpenAI(api_key=api_key)
                logger.info("OpenAI client initialized successfully")
            else:
                logger.warning("OpenAI client not available")
        except Exception as e:
            logger.error(f"Error initializing OpenAI client: {str(e)}")
    
    def generate_response(self, messages: List[Dict[str, str]], assistant_config: Dict[str, Any]) -> Dict[str, Any]:
        """Generate AI response"""
        try:
            if not self.client:
                return self.generate_demo_response(messages, assistant_config)
            
            # Prepare messages for OpenAI
            openai_messages = []
            
            # Add system message
            system_prompt = assistant_config.get("system_prompt", "You are a helpful AI assistant.")
            openai_messages.append({"role": "system", "content": system_prompt})
            
            # Add conversation messages
            for msg in messages:
                openai_messages.append({
                    "role": msg["role"],
                    "content": msg["content"]
                })
            
            # Get model configuration
            model = self.config.get("openai_model", "gpt-4")
            temperature = assistant_config.get("temperature", self.config.get("openai_temperature", 0.7))
            max_tokens = assistant_config.get("max_tokens", self.config.get("openai_max_tokens", 2000))
            
            start_time = time.time()
            
            # Make API call
            response = self.client.chat.completions.create(
                model=model,
                messages=openai_messages,
                temperature=temperature,
                max_tokens=max_tokens
            )
            
            response_time = time.time() - start_time
            
            # Extract response content
            content = response.choices[0].message.content
            
            # Calculate usage and cost
            usage = response.usage
            total_tokens = usage.total_tokens if usage else 0
            cost = self.calculate_cost(total_tokens, model)
            
            return {
                "content": content,
                "metadata": {
                    "model": model,
                    "total_tokens": total_tokens,
                    "prompt_tokens": usage.prompt_tokens if usage else 0,
                    "completion_tokens": usage.completion_tokens if usage else 0,
                    "cost": cost,
                    "response_time": response_time,
                    "demo_mode": False
                }
            }
            
        except Exception as e:
            logger.error(f"Error generating AI response: {str(e)}")
            return self.generate_demo_response(messages, assistant_config)
    
    def generate_demo_response(self, messages: List[Dict[str, str]], assistant_config: Dict[str, Any]) -> Dict[str, Any]:
        """Generate demo response when API is not available"""
        try:
            assistant_name = assistant_config.get("emoji", "🤖") + " " + "AI Assistant"
            
            # Get the last user message
            user_message = ""
            for msg in reversed(messages):
                if msg["role"] == "user":
                    user_message = msg["content"]
                    break
            
            # Generate contextual demo response
            demo_responses = [
                f"Thank you for your question about '{user_message[:50]}...'. As {assistant_name}, I would provide detailed insights and actionable recommendations based on my expertise in {', '.join(assistant_config.get('specialties', ['general consulting']))}.",
                
                f"I understand you're asking about '{user_message[:50]}...'. In my role as {assistant_name}, I would analyze this from multiple angles and provide strategic guidance tailored to your specific situation.",
                
                f"Great question about '{user_message[:50]}...'. As {assistant_name} with {assistant_config.get('expertise_level', 'professional')} experience, I would break this down into actionable steps and provide comprehensive recommendations.",
                
                f"I appreciate your inquiry regarding '{user_message[:50]}...'. Drawing from my expertise in {assistant_config.get('category', 'business consulting')}, I would provide detailed analysis and strategic insights to help you achieve your goals."
            ]
            
            import random
            content = random.choice(demo_responses)
            
            # Add demo-specific guidance
            content += f"\n\n**🎮 Demo Mode Active**\n\nTo get real AI responses with detailed analysis:\n1. Add your OpenAI API key to Streamlit secrets\n2. The assistant will then provide comprehensive, personalized guidance\n3. Full document analysis and context integration will be available\n\n**Current Configuration:**\n• Assistant: {assistant_name}\n• Expertise: {assistant_config.get('expertise_level', 'Professional')}\n• Specialties: {', '.join(assistant_config.get('specialties', ['General Consulting']))}"
            
            return {
                "content": content,
                "metadata": {
                    "model": "demo",
                    "total_tokens": len(content.split()),
                    "cost": 0.0,
                    "response_time": 0.5,
                    "demo_mode": True
                }
            }
            
        except Exception as e:
            logger.error(f"Error generating demo response: {str(e)}")
            return {
                "content": "I apologize, but I'm experiencing technical difficulties. Please try again or check your configuration.",
                "metadata": {
                    "model": "error",
                    "total_tokens": 0,
                    "cost": 0.0,
                    "response_time": 0.0,
                    "demo_mode": True,
                    "error": str(e)
                }
            }
    
    def calculate_cost(self, total_tokens: int, model: str) -> float:
        """Calculate cost based on token usage and model"""
        try:
            # Pricing per 1K tokens (approximate)
            pricing = {
                "gpt-4": {"input": 0.03, "output": 0.06},
                "gpt-4-turbo": {"input": 0.01, "output": 0.03},
                "gpt-3.5-turbo": {"input": 0.001, "output": 0.002}
            }
            
            # Use average pricing for simplicity
            if model in pricing:
                avg_price = (pricing[model]["input"] + pricing[model]["output"]) / 2
                return (total_tokens / 1000) * avg_price
            else:
                # Default pricing
                return (total_tokens / 1000) * 0.01
                
        except Exception as e:
            logger.error(f"Error calculating cost: {str(e)}")
            return 0.0

# Initialize AI chat manager
ai_chat_manager = AIChatManager(config_manager)

# ======================================================
# 🎨 UI COMPONENTS
# ======================================================

def render_header():
    """Render application header"""
    st.title(f"🚀 {config_manager.get('app_name')}")
    st.markdown(f"**{config_manager.get('app_description')}**")
    
    # Status indicators
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if config_manager.get("openai_api_key"):
            if ai_chat_manager.client:
                st.success("✅ OpenAI Connected")
            else:
                st.error("❌ OpenAI Failed")
        else:
            st.warning("🔑 API Key Required")
    
    with col2:
        if LANGCHAIN_CORE_AVAILABLE:
            st.success("✅ LangChain Ready")
        else:
            st.warning("⚠️ LangChain Limited")
    
    with col3:
        if FILE_PROCESSING_AVAILABLE:
            st.success("✅ File Processing Ready")
        else:
            st.warning("⚠️ Limited File Support")

def render_assistant_selector():
    """Render assistant selector"""
    st.sidebar.markdown("### 🤖 Professional AI Assistants")
    
    # Get all assistants
    all_assistants = assistant_profiles.get_all_assistants()
    
    # Category filter
    categories = list(set([assistant["category"] for assistant in all_assistants.values()]))
    selected_category = st.sidebar.selectbox(
        "Filter by Category", 
        ["All Categories"] + sorted(categories)
    )
    
    # Filter assistants
    filtered_assistants = {}
    for name, assistant in all_assistants.items():
        if selected_category == "All Categories" or assistant["category"] == selected_category:
            filtered_assistants[name] = assistant
    
    # Assistant selection
    assistant_names = list(filtered_assistants.keys())
    
    if not assistant_names:
        st.sidebar.warning("No assistants match the selected filters.")
        return None
    
    current_assistant = st.session_state.get("current_assistant", assistant_names[0])
    
    if current_assistant not in assistant_names:
        current_assistant = assistant_names[0]
    
    selected_assistant = st.sidebar.selectbox(
        "Select Your AI Assistant", 
        assistant_names, 
        index=assistant_names.index(current_assistant) if current_assistant in assistant_names else 0
    )
    
    # Update session state if changed
    if selected_assistant != st.session_state.get("current_assistant"):
        st.session_state.current_assistant = selected_assistant
        st.rerun()
    
    # Display assistant profile
    if selected_assistant in all_assistants:
        assistant_config = all_assistants[selected_assistant]
        
        st.sidebar.markdown(f"""
        <div style="
            background: linear-gradient(135deg, {config_manager.get('theme_primary_color')} 0%, {config_manager.get('theme_secondary_color')} 100%);
            padding: 20px;
            border-radius: 15px;
            color: white;
            margin: 15px 0;
            box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        ">
            <h3 style="margin: 0; color: white;">{assistant_config['emoji']} {selected_assistant}</h3>
            <p style="margin: 8px 0; opacity: 0.9; font-weight: bold;">{assistant_config['expertise_level']}</p>
            <p style="margin: 8px 0; opacity: 0.8;">{assistant_config['category']}</p>
            <p style="margin: 10px 0 0 0; opacity: 0.9; font-size: 0.9em;">{assistant_config['description']}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Detailed profile
        with st.sidebar.expander("📋 Assistant Details"):
            st.write(f"**🎯 Specialties:**")
            for specialty in assistant_config['specialties']:
                st.write(f"• {specialty}")
            
            st.write(f"\n**⚙️ Configuration:**")
            st.write(f"• Temperature: {assistant_config.get('temperature', 0.7)}")
            st.write(f"• Max Tokens: {assistant_config.get('max_tokens', 2000):,}")
        
        return selected_assistant
    
    return None

def render_file_upload():
    """Render file upload section"""
    st.sidebar.markdown("### 📁 File Management")
    
    # Upload configuration
    max_size_mb = config_manager.get("max_file_size_mb", 100)
    max_files = config_manager.get("max_files_per_session", 50)
    supported_types = config_manager.get("supported_file_types", [])
    
    st.sidebar.info(f"""
    **📊 Upload Limits**
    • Max file size: {max_size_mb}MB
    • Max files: {max_files}
    • Supported: {len(supported_types)} formats
    """)
    
    # File uploader
    uploaded_files = st.sidebar.file_uploader(
        "📤 Upload Files for Analysis",
        accept_multiple_files=True,
        type=supported_types if supported_types else None,
        help="Upload documents, data files, images, or code for AI analysis"
    )
    
    session_id = st.session_state.get("session_id", str(uuid.uuid4()))
    
    # Initialize session ID if not exists
    if "session_id" not in st.session_state:
        st.session_state.session_id = session_id
    
    # Process uploaded files
    if uploaded_files:
        # Check if files have changed
        file_signatures = [f"{f.name}_{f.size}" for f in uploaded_files]
        current_signatures = st.session_state.get("file_signatures", [])
        
        if file_signatures != current_signatures:
            st.session_state.file_signatures = file_signatures
            
            # Process files
            with st.sidebar.container():
                st.markdown("**🔄 Processing Files...**")
                
                processing_results = advanced_file_processor.process_uploaded_files(
                    uploaded_files, session_id
                )
                
                # Store results
                st.session_state.processing_results = processing_results
                
                # Display results
                if processing_results["status"] == "success":
                    successful_count = len(processing_results["processed_files"])
                    failed_count = len(processing_results["failed_files"])
                    
                    if successful_count > 0:
                        st.sidebar.success(f"✅ Processed {successful_count} files")
                    
                    if failed_count > 0:
                        st.sidebar.error(f"❌ Failed {failed_count} files")
                else:
                    st.sidebar.error(f"❌ Processing failed: {processing_results.get('message', 'Unknown error')}")
    
    # Display current files
    processing_results = st.session_state.get("processing_results")
    if processing_results and processing_results.get("processed_files"):
        st.sidebar.markdown("**📂 Current Files:**")
        
        processed_files = processing_results["processed_files"]
        
        # File summary
        col1, col2 = st.sidebar.columns(2)
        with col1:
            st.metric("Files", len(processed_files))
        with col2:
            total_size = sum(f.get("size", 0) for f in processed_files)
            st.metric("Size", f"{total_size / (1024*1024):.1f}MB")
        
        # File list
        with st.sidebar.expander("📋 File Details"):
            for file_info in processed_files:
                name = file_info.get("name", "Unknown")
                size = file_info.get("size", 0)
                method = file_info.get("processing_method", "unknown")
                
                st.write(f"**📄 {name}**")
                st.write(f"• Size: {size:,} bytes")
                st.write(f"• Method: {method}")
                st.write("---")
        
        # Clear files button
        if st.sidebar.button("🗑️ Clear All Files"):
            for key in ["processing_results", "file_signatures"]:
                if key in st.session_state:
                    del st.session_state[key]
            st.rerun()

def render_analytics():
    """Render analytics section"""
    st.sidebar.markdown("### 📊 Analytics")
    
    # Initialize analytics
    if "analytics" not in st.session_state:
        st.session_state.analytics = {
            "total_messages": 0,
            "total_tokens": 0,
            "total_cost": 0.0,
            "files_processed": 0
        }
    
    analytics = st.session_state.analytics
    
    # Display metrics
    col1, col2 = st.sidebar.columns(2)
    
    with col1:
        st.metric("Messages", analytics["total_messages"])
        st.metric("Files", analytics["files_processed"])
    
    with col2:
        st.metric("Tokens", f"{analytics['total_tokens']:,}")
        st.metric("Cost", f"${analytics['total_cost']:.4f}")

def generate_conversation_starters(assistant_config: Dict[str, Any], processing_results: Optional[Dict[str, Any]] = None) -> List[str]:
    """Generate contextual conversation starters"""
    try:
        category = assistant_config.get("category", "General")
        specialties = assistant_config.get("specialties", [])
        
        # Base suggestions by category
        suggestions_map = {
            "Business Strategy": [
                "Help me develop a strategic plan for market expansion",
                "Analyze my competitive landscape and positioning",
                "What are the key success factors for my industry?"
            ],
            "Marketing & Growth": [
                "Create a digital marketing strategy for customer acquisition",
                "How can I improve my conversion rates?",
                "What marketing channels should I prioritize?"
            ],
            "Finance & Investment": [
                "Analyze my financial performance and cash flow",
                "Help me create a budget and financial forecast",
                "What investment opportunities should I consider?"
            ],
            "Technology & Innovation": [
                "Develop a technology roadmap for digital transformation",
                "How can AI improve my business operations?",
                "What cybersecurity measures should I implement?"
            ]
        }
        
        # Get category-specific suggestions
        suggestions = suggestions_map.get(category, [
            "How can you help me with my business challenges?",
            "What insights can you provide about my industry?",
            "Help me analyze and improve my current situation"
        ])
        
        # Add file-specific suggestions if files are uploaded
        if processing_results and processing_results.get("processed_files"):
            file_suggestions = [
                "Analyze the uploaded files and provide key insights",
                "What patterns do you see in my data?",
                "Summarize the main findings from my documents"
            ]
            suggestions.extend(file_suggestions)
        
        return suggestions[:6]  # Limit to 6 suggestions
        
    except Exception as e:
        logger.error(f"Error generating conversation starters: {str(e)}")
        return [
            "How can you help me today?",
            "What are your key areas of expertise?",
            "Help me analyze my business situation"
        ]

def render_chat_interface():
    """Render main chat interface"""
    
    # Initialize session state
    if "messages" not in st.session_state:
        st.session_state.messages = {}
    if "current_assistant" not in st.session_state:
        st.session_state.current_assistant = "Strategic Business Consultant"
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if "analytics" not in st.session_state:
        st.session_state.analytics = {
            "total_messages": 0,
            "total_tokens": 0,
            "total_cost": 0.0,
            "files_processed": 0
        }
    
    # Render header and sidebar
    render_header()
    
    with st.sidebar:
        current_assistant = render_assistant_selector()
        render_file_upload()
        render_analytics()
        
        # Chat management
        st.markdown("### 🔧 Chat Management")
        
        if st.button("🗑️ Clear Chat"):
            if current_assistant in st.session_state.messages:
                del st.session_state.messages[current_assistant]
            st.rerun()
        
        if st.button("🔄 New Session"):
            st.session_state.session_id = str(uuid.uuid4())
            st.session_state.messages = {}
            if "processing_results" in st.session_state:
                del st.session_state["processing_results"]
            if "file_signatures" in st.session_state:
                del st.session_state["file_signatures"]
            st.rerun()
    
    # Main chat area
    if not current_assistant:
        st.error("Please select an AI assistant to continue.")
        return
    
    # Initialize messages for current assistant
    if current_assistant not in st.session_state.messages:
        st.session_state.messages[current_assistant] = []
    
    messages = st.session_state.messages[current_assistant]
    
    # Get assistant configuration
    all_assistants = assistant_profiles.get_all_assistants()
    assistant_config = all_assistants.get(current_assistant, {})
    
    # Display conversation or welcome screen
    if messages:
        st.markdown(f"### 💬 Conversation with {assistant_config.get('emoji', '🤖')} {current_assistant}")
        
        # Display messages
        for message in messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
                # Show metadata for assistant messages
                if message["role"] == "assistant" and "metadata" in message:
                    metadata = message["metadata"]
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.caption(f"💰 ${metadata.get('cost', 0):.4f}")
                    with col2:
                        st.caption(f"🔢 {metadata.get('total_tokens', 0):,}")
                    with col3:
                        st.caption(f"⏱️ {metadata.get('response_time', 0):.1f}s")
                    with col4:
                        if metadata.get('demo_mode'):
                            st.caption("🎮 Demo")
                        else:
                            st.caption("⚡ API")
    else:
        # Welcome screen
        st.markdown(f"### 💬 Welcome! Start a conversation with {assistant_config.get('emoji', '🤖')} {current_assistant}")
        
        # Assistant profile display
        if assistant_config:
            st.markdown(f"""
            <div style="
                background: linear-gradient(135deg, {config_manager.get('theme_primary_color')} 0%, {config_manager.get('theme_secondary_color')} 100%);
                padding: 30px;
                border-radius: 20px;
                color: white;
                margin: 20px 0;
                box-shadow: 0 8px 25px rgba(0,0,0,0.15);
            ">
                <h2 style="margin: 0; color: white; text-align: center;">{assistant_config.get('emoji', '🤖')} {current_assistant}</h2>
                <h4 style="margin: 10px 0; opacity: 0.9; text-align: center;">{assistant_config.get('expertise_level', 'AI Assistant')}</h4>
                <p style="margin: 15px 0; opacity: 0.9; text-align: center; font-size: 1.1em; line-height: 1.5;">{assistant_config.get('description', 'Professional AI assistant ready to help.')}</p>
                
                <div style="margin-top: 20px; text-align: center;">
                    <strong>🎯 Specialties:</strong><br>
                    {' • '.join(assistant_config.get('specialties', []))}
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        # Conversation starters
        st.markdown("**💡 Conversation Starters:**")
        
        suggestions = generate_conversation_starters(assistant_config, st.session_state.get("processing_results"))
        
        # Create suggestion buttons
        cols = st.columns(min(len(suggestions), 3))
        for i, suggestion in enumerate(suggestions):
            with cols[i % len(cols)]:
                if st.button(suggestion, key=f"suggestion_{i}"):
                    st.session_state.messages[current_assistant].append({
                        "role": "user",
                        "content": suggestion
                    })
                    st.rerun()
    
    # Show uploaded files summary
    processing_results = st.session_state.get("processing_results")
    if processing_results and processing_results.get("processed_files"):
        st.markdown("### 📁 Uploaded Files Ready for Analysis")
        
        processed_files = processing_results["processed_files"]
        
        # Create file display
        for file_info in processed_files[:3]:  # Show first 3 files
            name = file_info.get("name", "Unknown")
            size = file_info.get("size", 0)
            method = file_info.get("processing_method", "unknown")
            
            st.markdown(f"""
            <div style="
                border: 2px solid #28a745;
                border-radius: 12px;
                padding: 15px;
                margin: 8px 0;
                background: linear-gradient(145deg, #f8f9fa 0%, #e9ecef 100%);
            ">
                <div style="font-weight: bold; color: #333;">📄 {name}</div>
                <div style="font-size: 0.85em; color: #666;">
                    Size: {size:,} bytes | Method: {method}
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        if len(processed_files) > 3:
            st.info(f"... and {len(processed_files) - 3} more files ready for analysis")
    
    # Chat input
    if prompt := st.chat_input("Ask your AI assistant anything..."):
        # Add user message
        st.session_state.messages[current_assistant].append({
            "role": "user",
            "content": prompt
        })
        
        # Display user message
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate and display assistant response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                # Prepare context with file content if available
                context_messages = []
                
                # Add file context if available
                if processing_results and processing_results.get("processed_files"):
                    file_context = "=== UPLOADED FILES CONTEXT ===\n\n"
                    for file_info in processing_results["processed_files"]:
                        file_context += f"**File: {file_info.get('name', 'Unknown')}**\n"
                        file_context += f"{file_info.get('content', '')[:1000]}...\n\n"
                    
                    context_messages.append({
                        "role": "system",
                        "content": f"The user has uploaded files. Here's the content for context:\n\n{file_context}"
                    })
                
                # Add conversation history
                context_messages.extend(st.session_state.messages[current_assistant])
                
                # Generate response
                response_data = ai_chat_manager.generate_response(context_messages, assistant_config)
                
                # Display response
                st.markdown(response_data["content"])
                
                # Update analytics
                metadata = response_data["metadata"]
                st.session_state.analytics["total_messages"] += 1
                st.session_state.analytics["total_tokens"] += metadata.get("total_tokens", 0)
                st.session_state.analytics["total_cost"] += metadata.get("cost", 0.0)
                
                # Add assistant message to conversation
                st.session_state.messages[current_assistant].append({
                    "role": "assistant",
                    "content": response_data["content"],
                    "metadata": metadata
                })
        
        st.rerun()

# ======================================================
# 🚀 MAIN APPLICATION
# ======================================================

def main():
    """Main application function"""
    try:
        # Configure Streamlit page
        st.set_page_config(
            page_title=config_manager.get("app_name", "Advanced LangChain AI Assistant"),
            page_icon="🚀",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        # Custom CSS
        st.markdown(f"""
        <style>
        .stApp {{
            background: linear-gradient(135deg, {config_manager.get('theme_primary_color')}15 0%, {config_manager.get('theme_secondary_color')}15 100%);
        }}
        
        .stButton > button {{
            background: linear-gradient(135deg, {config_manager.get('theme_primary_color')} 0%, {config_manager.get('theme_secondary_color')} 100%);
            color: white;
            border: none;
            border-radius: 10px;
            padding: 0.5rem 1rem;
            font-weight: 600;
            transition: all 0.3s ease;
        }}
        
        .stButton > button:hover {{
            transform: translateY(-2px);
            box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        }}
        
        .stSelectbox > div > div {{
            background: white;
            border-radius: 10px;
        }}
        
        .stTextInput > div > div > input {{
            border-radius: 10px;
        }}
        
        .stFileUploader > div {{
            border-radius: 10px;
            border: 2px dashed {config_manager.get('theme_primary_color')};
        }}
        
        .stMetric {{
            background: white;
            padding: 1rem;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        
        .stExpander {{
            background: white;
            border-radius: 10px;
            border: 1px solid #e0e0e0;
        }}
        
        .stChatMessage {{
            background: white;
            border-radius: 15px;
            padding: 1rem;
            margin: 0.5rem 0;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        </style>
        """, unsafe_allow_html=True)
        
        # Render main chat interface
        render_chat_interface()
        
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        logger.error(f"Application error: {str(e)}")

if __name__ == "__main__":
    main()

